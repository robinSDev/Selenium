'''
Created on 11-Apr-2019
@author: Robin Singh
'''
from docx import Document
from docx.shared import Inches
import datetime
import pyautogui
import os
import time

class takeEvidences(object):
    '''
    classdocs
    '''
    global document

    def __init__(self, scriptName):
        '''
        Constructor
        '''
        runName = datetime.datetime.now().strftime("%I:%M%p_%B_%d_%Y ") + str(scriptName)
        self.document = Document()
        margin = Inches(0.5)
        #changing the page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
        
        self.document.add_heading(runName, 0)
    
    def addEvidence(self, title):
        time.sleep(0.15)
        picPath = os.getcwd() + '\\temp.png'
        self.document.add_paragraph(title)
        pic = pyautogui.screenshot()
        pic.save(picPath)
        self.document.add_picture(picPath, width=Inches(6.89), height=Inches(3.92))
    
    def __del__(self):
        self.document.save(os.getcwd() + '\\demo.docx')